import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import shutil

def create_college_api_doc():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    PRIMARY_COLOR = RGBColor(0, 51, 102)     # Dark Blue
    SECONDARY_COLOR = RGBColor(70, 130, 180)  # Steel Blue
    METHOD_POST_COLOR = RGBColor(40, 167, 69)
    METHOD_GET_COLOR = RGBColor(0, 123, 255)
    METHOD_PATCH_COLOR = RGBColor(255, 193, 7)
    METHOD_DELETE_COLOR = RGBColor(220, 53, 69)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CampXSync College Admin Module (Tenant Tier)\nREST APIs & Test Specification")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Technical Specification & Automated Test Execution Results")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()

    # Section 1: Executive Overview
    h1 = doc.add_heading("1. Executive & Architecture Overview", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph()
    p1.add_run(
        "The CampXSync College Admin Module (Tenant Tier) microservice (college-admin-service) provides full tenant-level "
        "administration capabilities across 5 core Epics: College Configuration, College Identity & Profile, College RBAC, "
        "College Reports & Analytics, and College Audit & Compliance. All 15 REST APIs defined in the specification have been "
        "implemented in Spring Boot 2.7.18 / Java 8 and 100% verified using Spring MockMvc tests."
    )

    # Overview Table
    t_over = doc.add_table(rows=6, cols=2)
    t_over.alignment = WD_TABLE_ALIGNMENT.CENTER
    over_data = [
        ("Microservice Name", "college-admin-service"),
        ("Base Package", "com.campsync.college"),
        ("Server Port", "8089"),
        ("Total User Stories Covered", "37 Stories across 5 Epics"),
        ("Total REST APIs", "15 Endpoints"),
        ("Automated Test Status", "17 / 17 Tests Passed (100% Pass Rate - BUILD SUCCESS)")
    ]

    for i, (k, v) in enumerate(over_data):
        c0 = t_over.cell(i, 0); c1 = t_over.cell(i, 1)
        c0.text = k; c1.text = v
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(10)
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")
        if k == "Automated Test Status":
            c1.paragraphs[0].runs[0].font.bold = True
            c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()

    # Modules and APIs data
    modules = [
        ("College Configuration Service", [
            ("API 01", "GET", "/v1/college-configs", "View institution configuration",
             "Returns active feature flags, branding, and operational settings for institution",
             "Stories 3, 6, 7", "Header: X-Institution-Id=inst-101",
             '[{"key":"attendance_module_enabled","value":true,"category":"feature_flag"},{"key":"theme_color","value":"#003366","category":"branding"}]',
             "200 OK - Config list returned"),
            ("API 02", "PATCH", "/v1/college-configs/{key}", "Update configuration setting",
             "Updates a setting (branding, feature flag, operational) for caller's institution",
             "Story 4", '{"value": "#0055A5"}',
             '{"key":"theme_color","value":"#0055A5","category":"branding","institutionId":"inst-101","updatedAt":"2026-08-10T10:35:53.808Z"}',
             "200 OK - Setting updated"),
            ("API 03", "GET", "/v1/college-configs/history", "View config change history",
             "Returns ordered audit list of prior values with timestamp and actor",
             "Story 5", "Header: X-Institution-Id=inst-101, Query: key=theme_color",
             '[{"key":"theme_color","previousValue":"#003366","newValue":"#0055A5","actor":"college-admin-1","timestamp":"2026-08-10T10:35:53.808Z"}]',
             "200 OK - History list returned")
        ]),
        ("College Identity & Profile Service", [
            ("API 04", "POST", "/v1/users", "Create user record",
             "Creates user of profile_type (student, faculty, staff, parent, alumni, admin) with status=active",
             "Stories 10, 15",
             '{"profileType":"student","name":"Jane Doe","email":"jane.doe@oxford.edu","profile":{"department":"CS","rollNo":"CS-2026-01"}}',
             '{"id":"usr-2cb936bd","institutionId":"inst-101","profileType":"student","name":"Jane Doe","status":"active"}',
             "201 Created - User provisioned"),
            ("API 05", "GET", "/v1/users/{id}", "View user identity & profile",
             "Fetches a single user record with profile sub-document",
             "Story 11", "Path Variable: id=usr-101, Header: X-Institution-Id=inst-101",
             '{"id":"usr-101","institutionId":"inst-101","profileType":"faculty","name":"Prof. Alan Turing","status":"active"}',
             "200 OK - User details returned"),
            ("API 06", "GET", "/v1/users", "List and filter users",
             "Filters users by profile_type and status, paginated",
             "Story 12", "Query: profile_type=faculty, status=active, page=0, size=10",
             '{"content":[{"id":"usr-101","name":"Prof. Alan Turing","profileType":"faculty"}],"page":0,"size":10,"totalElements":1}',
             "200 OK - Filtered list returned"),
            ("API 07", "PATCH", "/v1/users/{id}/status", "Transition account status",
             "Enforces state machine transitions (active<->suspended, active/suspended->deactivated)",
             "Stories 13, 16", '{"status": "suspended"}',
             '{"id":"usr-101","name":"Prof. Alan Turing","status":"suspended"}',
             "200 OK - Status updated"),
            ("API 08", "PATCH", "/v1/users/{id}/profile", "Update user profile sub-doc",
             "Updates type-specific profile sub-document fields",
             "Story 14", '{"profile": {"designation": "Senior Professor"}}',
             '{"id":"usr-101","profile":{"department":"Computer Science","designation":"Senior Professor"}}',
             "200 OK - Profile updated")
        ]),
        ("College RBAC Service", [
            ("API 09", "POST", "/v1/roles", "Define custom role",
             "Creates custom role for institution after validating permissions against global catalog",
             "Stories 18, 23",
             '{"name":"Department Head","description":"Head of Academic Department","permissions":["college:users:read","college:users:write"]}',
             '{"id":"role-3827880e","institutionId":"inst-101","name":"Department Head","permissions":["college:users:read"]}',
             "201 Created - Custom role defined"),
            ("API 10", "GET", "/v1/roles", "List custom roles",
             "Returns all custom roles defined for institution",
             "Story 18", "Header: X-Institution-Id=inst-101",
             '[{"id":"role-college-admin","name":"College Admin","permissions":["college:users:read","college:configs:write"]}]',
             "200 OK - Roles list returned"),
            ("API 11", "POST", "/v1/role-assignments", "Grant role to user",
             "Assigns role to user scoped to department/hostel/program",
             "Stories 19, 22", '{"userId":"usr-101","roleId":"role-college-admin","scope":"DEPARTMENT_CS"}',
             '{"id":"assign-8232a89d","institutionId":"inst-101","userId":"usr-101","roleId":"role-college-admin","scope":"DEPARTMENT_CS"}',
             "201 Created - Role granted"),
            ("API 12", "GET", "/v1/role-assignments/{user_id}/effective", "Resolve effective permissions",
             "Resolves tenant-scoped permissions, cached per session",
             "Stories 20, 36", "Path Variable: user_id=usr-101, Header: X-Institution-Id=inst-101",
             '{"userId":"usr-101","institutionId":"inst-101","roles":["College Admin"],"effectivePermissions":["college:users:read"]}',
             "200 OK - Effective permissions resolved"),
            ("API 13", "DELETE", "/v1/role-assignments/{id}", "Revoke role assignment",
             "Immediately revokes user role assignment",
             "Story 21", "Path Variable: id=assign-8232a89d",
             "No Content (HTTP 204)",
             "204 No Content - Role revoked")
        ]),
        ("College Reports & Analytics Service", [
            ("API 14", "GET", "/v1/college-analytics/dashboard", "View analytics dashboard",
             "Fetches precomputed rollups (enrollments, fee collections, exam results)",
             "Stories 25, 27, 28", "Query: metric=enrollment_summary, period=semester",
             '{"snapshotId":"tenant-snap-1786358153328","institutionId":"inst-101","data":{"totalEnrolledStudents":4850}}',
             "200 OK - Dashboard analytics returned"),
            ("API 15", "POST", "/v1/college-analytics/recompute", "Trigger analytics recompute",
             "Triggers asynchronous recompute of rollups",
             "Story 26", '{"metric":"enrollment_summary","period":"semester"}',
             '{"jobId":"tenant-job-1786358153774","institutionId":"inst-101","status":"ACCEPTED"}',
             "202 Accepted - Recompute queued")
        ]),
        ("College Audit & Compliance Service", [
            ("API 16", "GET", "/v1/audit-logs", "Query institution audit trail",
             "Returns paginated audit log entries implicitly scoped to institution_id",
             "Stories 30, 31, 32, 33", "Query: page=0, size=5",
             '{"content":[{"id":"tenant-audit-1","eventType":"UserCreated","sourceModule":"College Identity Service"}],"totalElements":3}',
             "200 OK - Audit trail returned")
        ])
    ]

    for sec_idx, (mod_name, apis) in enumerate(modules, start=2):
        h_sec = doc.add_heading(f"{sec_idx}. {mod_name}", level=1)
        h_sec.runs[0].font.color.rgb = PRIMARY_COLOR

        for api_num, method, path, title, desc, story, req, resp, status in apis:
            p_api = doc.add_paragraph()
            r_num = p_api.add_run(f"• {api_num}: ")
            r_num.font.bold = True

            r_mth = p_api.add_run(f"[{method}] ")
            r_mth.font.bold = True
            if method == "POST": r_mth.font.color.rgb = METHOD_POST_COLOR
            elif method == "GET": r_mth.font.color.rgb = METHOD_GET_COLOR
            elif method == "PATCH": r_mth.font.color.rgb = METHOD_PATCH_COLOR
            elif method == "DELETE": r_mth.font.color.rgb = METHOD_DELETE_COLOR

            r_path = p_api.add_run(f"{path} — {title}")
            r_path.font.bold = True
            r_path.font.size = Pt(11)

            t_api = doc.add_table(rows=5, cols=2)
            t_api.alignment = WD_TABLE_ALIGNMENT.CENTER

            rows_info = [
                ("User Story Covered", story),
                ("Description & Logic", desc),
                ("Sample Request Input", req),
                ("Sample JSON Response", resp),
                ("Test Verification Status", f"PASSED — {status}")
            ]

            for r_idx, (lbl, val) in enumerate(rows_info):
                c0 = t_api.cell(r_idx, 0); c1 = t_api.cell(r_idx, 1)
                c0.text = lbl; c1.text = val
                c0.paragraphs[0].runs[0].font.bold = True
                set_cell_background(c0, "F0F4F8")
                set_cell_background(c1, "FAFAFA")
                if lbl == "Test Verification Status":
                    c1.paragraphs[0].runs[0].font.bold = True
                    c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    set_cell_background(c1, "E8F5E9")

            for row in t_api.rows:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9.5)

            doc.add_paragraph()

    output_path = r"d:\CampXSync\CampSync_College_Admin_Module_APIs_Documentation.docx"
    doc.save(output_path)
    shutil.copy(output_path, r"d:\CampXSync\CampXSync\CampSync_College_Admin_Module_APIs_Documentation.docx")
    shutil.copy(output_path, r"d:\CampXSync\Documentation\AdminModule\CampSync_College_Admin_Module_APIs_Documentation.docx")

    v1_path = r"d:\CampXSync\CampSync_College_Admin_Module_APIs_Documentation_v1.0.docx"
    doc.save(v1_path)
    shutil.copy(v1_path, r"d:\CampXSync\CampXSync\CampSync_College_Admin_Module_APIs_Documentation_v1.0.docx")
    shutil.copy(v1_path, r"d:\CampXSync\Documentation\AdminModule\CampSync_College_Admin_Module_APIs_Documentation_v1.0.docx")

    print("College Admin Module APIs Documentation Word Document generated successfully!")

if __name__ == "__main__":
    create_college_api_doc()
