import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import shutil

def create_gateway_user_guide_doc():
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Theme Colors
    PRIMARY_COLOR = RGBColor(0, 51, 102)     # Dark Blue
    SECONDARY_COLOR = RGBColor(70, 130, 180)  # Steel Blue
    SUCCESS_COLOR = RGBColor(0, 128, 0)      # Green

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CampXSync API Gateway User Guide & Technical Documentation")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Single Ingress Entrypoint (Port 8080), Security, JWT Authentication & Reverse Proxy Routing")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()

    # 1. Architectural Overview
    h1 = doc.add_heading("1. Architectural Overview", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph()
    p1.add_run(
        "The CampXSync API Gateway (api-gateway) serves as the centralized ingress controller and reverse proxy running on Port 8080. "
        "All client applications—including Web Portals, Mobile Apps, and External Integrations—communicate exclusively through the Gateway. "
        "The Gateway handles authentication token validation (JWT), MDC trace context injection, access logging, CORS preflight, "
        "and dynamic request dispatching to downstream microservices."
    )

    t_meta = doc.add_table(rows=6, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Gateway Service Name", "api-gateway"),
        ("Ingress Port", "8080"),
        ("Downstream Target 1", "Platform Admin Service (http://localhost:8088)"),
        ("Downstream Target 2", "College Admin Service (http://localhost:8089)"),
        ("Authentication Mechanism", "JWT Bearer Token (logger.jwt.JwtProvider) / Ingress Headers"),
        ("Logging Framework", "com.campxsync:logger:1.0.0 (AppLogger & AuditContextHolder)")
    ]
    for i, (k, v) in enumerate(meta_info):
        c0 = t_meta.cell(i, 0); c1 = t_meta.cell(i, 1)
        c0.text = k; c1.text = v
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9.5)
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")

    doc.add_paragraph()

    # 2. How to Start the Services
    h2 = doc.add_heading("2. Step-by-Step Local Startup Guide", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("To run the full CampXSync microservices ecosystem locally, start each service in a separate terminal:")

    steps = [
        ("Terminal 1: Platform Admin Microservice (Port 8088)",
         "cd d:\\CampXSync\\CampXSync\\services\\platform-admin-service\nmvn spring-boot:run"),
        ("Terminal 2: College Admin Microservice (Port 8089)",
         "cd d:\\CampXSync\\CampXSync\\services\\college-admin-service\nmvn spring-boot:run"),
        ("Terminal 3: Central API Gateway Microservice (Port 8080)",
         "cd d:\\CampXSync\\CampXSync\\api-gateway\nmvn spring-boot:run")
    ]

    for title, cmd in steps:
        p_step = doc.add_paragraph()
        r_step = p_step.add_run(f"• {title}")
        r_step.font.bold = True
        r_step.font.size = Pt(10.5)

        p_cmd = doc.add_paragraph()
        r_cmd = p_cmd.add_run(cmd)
        r_cmd.font.name = "Consolas"
        r_cmd.font.size = Pt(9.5)

    doc.add_paragraph()

    # 3. Gateway Route Dispatch Matrix
    h3 = doc.add_heading("3. API Gateway Routing & Dispatch Matrix", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("All requests sent to http://localhost:8080 are dynamically routed according to the endpoint prefix:")

    routes = [
        ("/v1/institutes/**", "Platform Admin Service", "http://localhost:8088/v1/institutes/**", "Institute Provisioning & Lifecycle"),
        ("/v1/platform-configs/**", "Platform Admin Service", "http://localhost:8088/v1/platform-configs/**", "Global Platform Configurations"),
        ("/v1/platform-roles/**", "Platform Admin Service", "http://localhost:8088/v1/platform-roles/**", "Platform Super Admin RBAC"),
        ("/v1/platform-role-assignments/**", "Platform Admin Service", "http://localhost:8088/v1/platform-role-assignments/**", "Platform Staff Role Grants"),
        ("/v1/policies/**", "Platform Admin Service", "http://localhost:8088/v1/policies/**", "Data Governance & Retention Policies"),
        ("/v1/billing-accounts/**", "Platform Admin Service", "http://localhost:8088/v1/billing-accounts/**", "Subscriptions & Charge Settlement"),
        ("/v1/analytics/**", "Platform Admin Service", "http://localhost:8088/v1/analytics/**", "Cross-Tenant Platform Rollups"),
        ("/v1/compliance-checks/**", "Platform Admin Service", "http://localhost:8088/v1/compliance-checks/**", "Automated Policy Audits"),
        ("/v1/platform-audit-logs/**", "Platform Admin Service", "http://localhost:8088/v1/platform-audit-logs/**", "Platform Master Audit Trail"),
        ("/v1/college-configs/**", "College Admin Service", "http://localhost:8089/v1/college-configs/**", "Tenant Feature Flags & Branding"),
        ("/v1/users/**", "College Admin Service", "http://localhost:8089/v1/users/**", "Student/Faculty Identity & Profiles"),
        ("/v1/roles/**", "College Admin Service", "http://localhost:8089/v1/roles/**", "Tenant Custom Roles"),
        ("/v1/role-assignments/**", "College Admin Service", "http://localhost:8089/v1/role-assignments/**", "Tenant User Role Assignments"),
        ("/v1/college-analytics/**", "College Admin Service", "http://localhost:8089/v1/college-analytics/**", "Institution Operational Metrics"),
        ("/v1/audit-logs/**", "College Admin Service", "http://localhost:8089/v1/audit-logs/**", "Tenant Isolated Audit Trail")
    ]

    t_route = doc.add_table(rows=len(routes) + 1, cols=4)
    t_route.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Gateway Request Prefix (Port 8080)", "Destination Service", "Routed Downstream URL", "Module Description"]
    for j, h_text in enumerate(headers):
        cell = t_route.cell(0, j)
        cell.text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (prefix, svc, target_url, desc) in enumerate(routes, start=1):
        c0 = t_route.cell(i, 0); c1 = t_route.cell(i, 1); c2 = t_route.cell(i, 2); c3 = t_route.cell(i, 3)
        c0.text = prefix; c1.text = svc; c2.text = target_url; c3.text = desc
        bg = "F9FBFD" if i % 2 == 0 else "FFFFFF"
        for c in [c0, c1, c2, c3]:
            c.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_background(c, bg)

    doc.add_paragraph()

    # 4. Sample cURL & API Usage Examples
    h4 = doc.add_heading("4. Sample cURL Request & Verification Examples", level=1)
    h4.runs[0].font.color.rgb = PRIMARY_COLOR

    examples = [
        ("Example 1: Provision Institute (Platform Admin Service via Gateway)",
         "curl -X POST http://localhost:8080/v1/institutes \\\n"
         "  -H \"Content-Type: application/json\" \\\n"
         "  -H \"X-User-Id: admin-super-1\" \\\n"
         "  -d '{\n"
         "    \"name\": \"Harvard University\",\n"
         "    \"subdomain\": \"harvard\",\n"
         "    \"planId\": \"plan-enterprise\",\n"
         "    \"tenancyTier\": \"DEDICATED_TENANT\"\n"
         "  }'",
         "HTTP/1.1 201 Created\nX-Trace-Id: 91e26c2911914a788df517857e1c08dd\nX-Response-Time: 38ms\n\n"
         "{\n  \"id\": \"inst-4a0a47ba\",\n  \"name\": \"Harvard University\",\n  \"subdomain\": \"harvard\",\n  \"status\": \"onboarding\"\n}"),

        ("Example 2: Create College User (College Admin Service via Gateway)",
         "curl -X POST http://localhost:8080/v1/users \\\n"
         "  -H \"Content-Type: application/json\" \\\n"
         "  -H \"X-Institution-Id: inst-101\" \\\n"
         "  -H \"X-User-Id: admin-college-1\" \\\n"
         "  -d '{\n"
         "    \"profileType\": \"student\",\n"
         "    \"name\": \"Alice Smith\",\n"
         "    \"email\": \"alice.smith@oxford.edu\",\n"
         "    \"profile\": {\"department\": \"Computer Science\"}\n"
         "  }'",
         "HTTP/1.1 201 Created\nX-Trace-Id: 93ffcfc8776a43a9962b2c889333f648\nX-Response-Time: 24ms\n\n"
         "{\n  \"id\": \"usr-1fa364ea\",\n  \"institutionId\": \"inst-101\",\n  \"name\": \"Alice Smith\",\n  \"status\": \"active\"\n}"),

        ("Example 3: Health Actuator Status Check",
         "curl -X GET http://localhost:8080/actuator/health",
         "HTTP/1.1 200 OK\n\n{\"status\": \"UP\", \"components\": {\"ping\": {\"status\": \"UP\"}}}")
    ]

    for title, req_code, resp_code in examples:
        p_ex = doc.add_paragraph()
        r_ex = p_ex.add_run(f"• {title}")
        r_ex.font.bold = True
        r_ex.font.size = Pt(10.5)

        t_ex = doc.add_table(rows=2, cols=2)
        t_ex.alignment = WD_TABLE_ALIGNMENT.CENTER

        c00 = t_ex.cell(0, 0); c01 = t_ex.cell(0, 1)
        c00.text = "cURL Request"; c01.text = req_code
        c00.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c00, "F0F4F8"); set_cell_background(c01, "FAFAFA")

        c10 = t_ex.cell(1, 0); c11 = t_ex.cell(1, 1)
        c10.text = "Expected Gateway Output"; c11.text = resp_code
        c10.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c10, "F0F4F8"); set_cell_background(c11, "FAFAFA")

        for row in t_ex.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)

        doc.add_paragraph()

    output_path = r"d:\CampXSync\CampXSync_API_Gateway_User_Guide_and_Documentation.docx"
    doc.save(output_path)

    for dst_dir in [r"d:\CampXSync\CampXSync", r"d:\CampXSync\Documentation\AdminModule"]:
        try:
            shutil.copy(output_path, dst_dir)
        except Exception:
            pass

    print(f"API Gateway User Guide Document generated successfully at {output_path}!")

if __name__ == "__main__":
    create_gateway_user_guide_doc()
