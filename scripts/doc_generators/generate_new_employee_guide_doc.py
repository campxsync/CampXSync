import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import shutil

def create_new_employee_guide_doc():
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Theme Colors
    PRIMARY_COLOR = RGBColor(0, 51, 102)     # Dark Blue
    SECONDARY_COLOR = RGBColor(70, 130, 180)  # Steel Blue
    SUCCESS_COLOR = RGBColor(0, 128, 0)      # Green

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CampXSync API Gateway & Microservices Developer Onboarding Guide")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Engineering Reference & Practical Training Example for New Developers")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()

    # 1. Welcome & Architecture Overview
    h1 = doc.add_heading("1. Welcome & Ecosystem Architecture", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph()
    p1.add_run(
        "Welcome to the CampXSync Engineering Team! This document serves as a practical, hands-on onboarding guide "
        "for understanding the CampXSync microservices architecture, API Gateway routing, shared logging standards, and API usage."
    )

    t_arch = doc.add_table(rows=5, cols=3)
    t_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Service Name", "Port", "Core Responsibility"]
    for j, h_text in enumerate(headers):
        cell = t_arch.cell(0, j)
        cell.text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    svc_rows = [
        ("api-gateway", "8080", "Single ingress proxy, JWT validation, MDC trace context injection, CORS, and reverse routing"),
        ("platform-admin-service", "8088", "Platform Tier: Institutes onboarding, global configs, platform RBAC, data governance, billing"),
        ("college-admin-service", "8089", "Tenant Tier: College configs, student/faculty profiles, tenant RBAC, analytics, tenant audit"),
        ("logger (shared lib)", "N/A", "Shared Maven library providing AppLogger, AuditLogger, JwtProvider, TtlCache, Encryption")
    ]
    for i, (svc, port, desc) in enumerate(svc_rows, start=1):
        c0 = t_arch.cell(i, 0); c1 = t_arch.cell(i, 1); c2 = t_arch.cell(i, 2)
        c0.text = svc; c1.text = port; c2.text = desc
        bg = "F9FBFD" if i % 2 == 0 else "FFFFFF"
        for c in [c0, c1, c2]:
            c.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_background(c, bg)

    doc.add_paragraph()

    # 2. Local Environment Setup
    h2 = doc.add_heading("2. Local Developer Environment Setup", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("Follow these exact steps to build and launch the local environment:")

    setup_steps = [
        ("Step 1: Install Shared Logger to Local Maven Repo",
         "cd d:\\CampXSync\\CampXSync\\logger\nmvn clean install"),
        ("Step 2: Start Platform Admin Microservice (Port 8088)",
         "cd d:\\CampXSync\\CampXSync\\services\\platform-admin-service\nmvn spring-boot:run"),
        ("Step 3: Start College Admin Microservice (Port 8089)",
         "cd d:\\CampXSync\\CampXSync\\services\\college-admin-service\nmvn spring-boot:run"),
        ("Step 4: Start API Gateway Microservice (Port 8080)",
         "cd d:\\CampXSync\\CampXSync\\api-gateway\nmvn spring-boot:run")
    ]

    for title, cmd in setup_steps:
        p_s = doc.add_paragraph()
        r_s = p_s.add_run(f"• {title}")
        r_s.font.bold = True
        r_s.font.size = Pt(10)

        p_c = doc.add_paragraph()
        r_c = p_c.add_run(cmd)
        r_c.font.name = "Consolas"
        r_c.font.size = Pt(9)

    doc.add_paragraph()

    # 3. API Gateway Routing Matrix
    h3 = doc.add_heading("3. API Gateway Routing Matrix (Port 8080)", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("All API requests must target http://localhost:8080. The Gateway dispatches requests based on path prefixes:")

    routes = [
        ("/v1/institutes/**", "Platform Admin Service (8088)", "Institute onboarding & lifecycle state machine"),
        ("/v1/platform-configs/**", "Platform Admin Service (8088)", "Global platform configuration toggles & audit history"),
        ("/v1/platform-roles/**", "Platform Admin Service (8088)", "Super Admin roles & staff role assignments"),
        ("/v1/policies/**", "Platform Admin Service (8088)", "Data governance, retention & residency rules"),
        ("/v1/billing-accounts/**", "Platform Admin Service (8088)", "Subscriptions, plan changes & payment settlement"),
        ("/v1/analytics/**", "Platform Admin Service (8088)", "Platform-wide analytics rollups & recompute jobs"),
        ("/v1/compliance-checks/**", "Platform Admin Service (8088)", "Automated compliance checks & non-compliant flags"),
        ("/v1/platform-audit-logs/**", "Platform Admin Service (8088)", "Master platform audit log querying"),
        ("/v1/college-configs/**", "College Admin Service (8089)", "Tenant branding, theme color & feature flags"),
        ("/v1/users/**", "College Admin Service (8089)", "Student/Faculty identity profiles & status transitions"),
        ("/v1/roles/**", "College Admin Service (8089)", "Tenant custom roles & user role assignments"),
        ("/v1/college-analytics/**", "College Admin Service (8089)", "Tenant operational dashboard & snapshot recompute"),
        ("/v1/audit-logs/**", "College Admin Service (8089)", "Tenant isolated audit trail querying")
    ]

    t_r = doc.add_table(rows=len(routes) + 1, cols=3)
    t_r.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_headers = ["Ingress Prefix (Port 8080)", "Target Downstream Service", "Functional Area"]
    for j, h_text in enumerate(r_headers):
        cell = t_r.cell(0, j)
        cell.text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (prefix, target, area) in enumerate(routes, start=1):
        c0 = t_r.cell(i, 0); c1 = t_r.cell(i, 1); c2 = t_r.cell(i, 2)
        c0.text = prefix; c1.text = target; c2.text = area
        bg = "F9FBFD" if i % 2 == 0 else "FFFFFF"
        for c in [c0, c1, c2]:
            c.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_background(c, bg)

    doc.add_paragraph()

    # 4. Code Pattern Examples for New Developers
    h4 = doc.add_heading("4. Code Standards & Development Examples", level=1)
    h4.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("All new service classes must integrate AppLogger and AuditLogger as demonstrated below:")

    code_snip = (
        "package com.campsync.college.service.impl;\n\n"
        "import logger.constants.AuditConstants;\n"
        "import logger.logging.AppLogger;\n"
        "import logger.logging.AuditLogger;\n"
        "import org.springframework.stereotype.Service;\n\n"
        "@Service\n"
        "public class SampleServiceImpl implements SampleService {\n"
        "    private static final AppLogger log = AppLogger.getLogger(SampleServiceImpl.class);\n\n"
        "    public void processOrder(String orderId) {\n"
        "        log.info(\"Processing request for orderId: {}\", orderId);\n\n"
        "        AuditLogger.builder()\n"
        "                .action(AuditConstants.ACTION_CREATE)\n"
        "                .entity(\"ORDER\", orderId)\n"
        "                .success()\n"
        "                .message(\"Order processed successfully\")\n"
        "                .log();\n"
        "    }\n"
        "}"
    )

    p_code = doc.add_paragraph()
    r_code = p_code.add_run(code_snip)
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(8.5)

    doc.add_paragraph()

    # 5. Developer Troubleshooting Reference
    h5 = doc.add_heading("5. Troubleshooting & FAQ", level=1)
    h5.runs[0].font.color.rgb = PRIMARY_COLOR

    faq_items = [
        ("Q: Receiving HTTP 502 Bad Gateway from Port 8080?",
         "Check if downstream microservices (8088 or 8089) are running. Start them using `mvn spring-boot:run` in their respective directories."),
        ("Q: Receiving HTTP 401 Unauthorized?",
         "Verify that your JWT token is signed with the secret key 'super_secret_signing_key_for_campxsync_platform_2026' or pass 'X-User-Id' header for testing."),
        ("Q: Maven compilation error 'cannot find symbol AppLogger'?",
         "Re-run `mvn clean install` inside `d:\\CampXSync\\CampXSync\\logger` to update the local Maven repository artifact (`com.campxsync:logger:1.0.0`).")
    ]

    for q, a in faq_items:
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(q)
        r_q.font.bold = True
        r_q.font.size = Pt(10)

        p_a = doc.add_paragraph()
        r_a = p_a.add_run(a)
        r_a.font.size = Pt(9.5)

    output_path = r"d:\CampXSync\CampXSync_API_Gateway_New_Employee_Onboarding_Guide.docx"
    doc.save(output_path)

    for dst_dir in [r"d:\CampXSync\CampXSync", r"d:\CampXSync\Documentation\AdminModule"]:
        try:
            shutil.copy(output_path, dst_dir)
        except Exception:
            pass

    print(f"New Employee Onboarding Guide generated successfully at {output_path}!")

if __name__ == "__main__":
    create_new_employee_guide_doc()
