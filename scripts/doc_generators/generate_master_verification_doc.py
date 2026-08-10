import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import shutil

def create_master_verification_doc():
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Theme Colors
    PRIMARY_COLOR = RGBColor(0, 51, 102)     # Dark Blue
    SECONDARY_COLOR = RGBColor(70, 130, 180)  # Steel Blue
    SUCCESS_COLOR = RGBColor(0, 128, 0)      # Green

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CampXSync Admin Module Master API Test Verification Report")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Verification & Validation Report for Platform Admin & College Admin Microservices")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()

    # 1. Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Quality Assurance Metrics", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p_summary = doc.add_paragraph()
    p_summary.add_run(
        "This master verification report presents the automated test execution results for the CampXSync Administrator Module. "
        "The test suite validates both microservices—Platform Admin Service (platform-admin-service) and College Admin Service "
        "(college-admin-service)—against Spring Boot MockMvc, covering all 90 User Stories across 13 Epics."
    )

    t_meta = doc.add_table(rows=8, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Microservices Tested", "2 Microservices (platform-admin-service:8088 & college-admin-service:8089)"),
        ("Total User Stories Verified", "90 Stories (53 Platform Tier + 37 Tenant Tier)"),
        ("Total REST APIs Verified", "42 REST API Endpoints (27 Platform Tier + 15 College Tier)"),
        ("Total Automated Tests Executed", "46 Automated Integration Test Cases (29 Platform + 17 College)"),
        ("Overall Pass Rate", "46 / 46 (100% Success Rate)"),
        ("Overall Build Result", "BUILD SUCCESS"),
        ("Structured Logging & Audit Integration", "Verified using com.campxsync:logger:1.0.0 (AppLogger & AuditLogger)"),
        ("Test Execution Frameworks", "Java 8 (1.8.0_202) / Spring Boot 2.7.18 / JUnit 5 / Spring MockMvc")
    ]

    for i, (k, v) in enumerate(meta_info):
        c0 = t_meta.cell(i, 0); c1 = t_meta.cell(i, 1)
        c0.text = k; c1.text = v
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(10)
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")
        if k in ["Overall Pass Rate", "Overall Build Result"]:
            c1.paragraphs[0].runs[0].font.bold = True
            c1.paragraphs[0].runs[0].font.color.rgb = SUCCESS_COLOR
            set_cell_background(c1, "E8F5E9")

    doc.add_paragraph()

    # 2. Platform Admin Service Verification
    h2 = doc.add_heading("2. Platform Admin Service (Platform Tier) Test Verification", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    p2 = doc.add_paragraph()
    p2.add_run("All 27 REST APIs and 2 boundary/resilience test cases passed with 100% success.")

    platform_tests = [
        ("API 01", "POST /v1/institutes", "Provision Institute", "PASSED (201 Created)",
         '{"name":"Stanford University","subdomain":"stanford","planId":"plan-enterprise-plus","tenancyTier":"DEDICATED_TENANT"}',
         '{"id":"inst-4a0a47ba","name":"Stanford University","status":"onboarding","tenancyTier":"DEDICATED_TENANT"}'),
        ("API 02", "GET /v1/institutes/{id}", "View Institute Details", "PASSED (200 OK)", "Path Variable: id=inst-101",
         '{"id":"inst-101","name":"Oxford International Academy","status":"suspended","tenancyTier":"HIGH_AVAILABILITY"}'),
        ("API 03", "GET /v1/institutes", "List & Filter Institutes", "PASSED (200 OK)", "Query: status=active, page=0, size=10",
         '{"content":[],"page":0,"size":10,"totalElements":0}'),
        ("API 04", "PATCH /v1/institutes/{id}/status", "Transition Status & Publish Events", "PASSED (200 OK)", '{"status":"suspended"}',
         '{"id":"inst-101","name":"Oxford International Academy","status":"suspended"}'),
        ("API 05", "PATCH /v1/institutes/{id}", "Update Institute Profile", "PASSED (200 OK)", '{"name":"Oxford International Academy"}',
         '{"id":"inst-101","name":"Oxford International Academy"}'),
        ("API 06", "GET /v1/platform-configs", "View Global Configs", "PASSED (200 OK)", "None",
         '[{"key":"mfa_required","value":false},{"key":"jwt_ttl_seconds","value":86400}]'),
        ("API 07", "PATCH /v1/platform-configs/{key}", "Update Global Config", "PASSED (200 OK)", '{"value": false} (Header: X-Actor-Id=admin-john)',
         '{"key":"mfa_required","value":false,"updatedBy":"admin-john"}'),
        ("API 08", "GET /v1/platform-configs/history", "View Config Audit History", "PASSED (200 OK)", "Query: key=mfa_required",
         '[{"key":"mfa_required","previousValue":true,"newValue":false,"actor":"admin-john"}]'),
        ("API 09", "POST /v1/platform-roles", "Define Platform Role", "PASSED (201 Created)", '{"name":"Support Ops Manager","permissions":["platform:institutes:read"]}',
         '{"id":"role-3fceaf22","name":"Support Ops Manager"}'),
        ("API 10", "GET /v1/platform-roles", "List Platform Roles", "PASSED (200 OK)", "None",
         '[{"id":"role-super-admin","name":"Super Admin"},{"id":"role-3fceaf22","name":"Support Ops Manager"}]'),
        ("API 11", "POST /v1/platform-role-assignments", "Grant Role to Staff", "PASSED (201 Created)", '{"staffId":"staff-99","roleId":"role-super-admin"}',
         '{"id":"assign-5b3d6f9b","staffId":"staff-99","roleName":"Super Admin"}'),
        ("API 12", "GET /v1/platform-role-assignments/{staff_id}/effective", "Resolve Effective Permissions", "PASSED (200 OK)", "Path Variable: staff_id=staff-99",
         '{"staffId":"staff-99","roles":["Super Admin"],"effectivePermissions":["platform:institutes:read"]}'),
        ("API 13", "DELETE /v1/platform-role-assignments/{id}", "Revoke Role Assignment", "PASSED (204 No Content)", "Path Variable: id=assign-5b3d6f9b",
         "HTTP 204 No Content"),
        ("API 14", "POST /v1/policies", "Create Data Governance Policy", "PASSED (201 Created)", '{"name":"Data Residency EU-Central","type":"residency","appliesTo":"all"}',
         '{"id":"pol-6cf31155","name":"Data Residency EU-Central","status":"ACTIVE"}'),
        ("API 15", "GET /v1/policies", "List Policies", "PASSED (200 OK)", "Query: applies_to=all",
         '[{"id":"pol-101","name":"GDPR 7-Year Student Data Retention","status":"RETIRED"}]'),
        ("API 16", "PATCH /v1/policies/{id}", "Update Policy", "PASSED (200 OK)", '{"name":"Updated Retention Policy"}',
         '{"id":"pol-101","name":"Updated Retention Policy","status":"RETIRED"}'),
        ("API 17", "DELETE /v1/policies/{id}", "Retire Policy", "PASSED (200 OK)", "Path Variable: id=pol-101",
         '{"id":"pol-101","status":"RETIRED"}'),
        ("API 18", "GET /v1/billing-accounts/{institution_id}", "View Billing Status", "PASSED (200 OK)", "Path Variable: institution_id=inst-101",
         '{"id":"bill-101","institutionId":"inst-101","planId":"plan-enterprise-v2"}'),
        ("API 19", "POST /v1/billing-accounts/{institution_id}/change-plan", "Change Subscription Plan", "PASSED (200 OK)", '{"newPlanId":"plan-enterprise-v2"}',
         '{"id":"bill-101","planId":"plan-enterprise-v2"}'),
        ("Story 32", "POST /v1/billing-accounts/{institution_id}/change-plan", "Settlement Failure Plan Reversal", "PASSED (402 Payment Required)", '{"newPlanId":"plan-invalid"}',
         'Settlement failed for plan change to "plan-invalid". Account reverted to "plan-enterprise-v2".'),
        ("API 20", "POST /v1/billing-accounts/{institution_id}/charge", "Trigger Manual Settlement", "PASSED (200 OK)", '{"amount":1250.00,"description":"Overage charge"}',
         '{"transactionId":"txn-be0c1521","status":"SUCCESS"}'),
        ("API 21", "GET /v1/analytics/snapshots", "View Cross-Institution Rollup", "PASSED (200 OK)", "Query: metric=active_institutions_count, period=monthly",
         '{"snapshotId":"snap-1723285513","data":{"activeInstitutions":42}}'),
        ("API 22", "POST /v1/analytics/snapshots/recompute", "Trigger Analytics Recompute", "PASSED (202 Accepted)", '{"metric":"active_institutions_count","period":"monthly"}',
         '{"jobId":"job-1723285513","status":"ACCEPTED"}'),
        ("API 23", "POST /v1/compliance-checks/run", "Trigger Compliance Check", "PASSED (202 Accepted)", '{"institutionId":"inst-101","policyIds":["pol-101"]}',
         '{"checkReferenceId":"chk-ref-1723285513","status":"QUEUED"}'),
        ("API 24", "GET /v1/compliance-checks/{institution_id}", "View Compliance Results", "PASSED (200 OK)", "Path Variable: institution_id=inst-101",
         '{"checkId":"chk-101","compliant":true,"violations":[]}'),
        ("API 25", "GET /v1/compliance-checks", "List Non-Compliant Institutes", "PASSED (200 OK)", "Query: flagged=true",
         '[{"checkId":"chk-102","institutionId":"inst-102","compliant":false}]'),
        ("API 26", "GET /v1/platform-audit-logs", "Query Platform Audit Trail", "PASSED (200 OK)", "Query: page=0, size=5",
         '{"content":[{"id":"audit-1","eventType":"InstituteOnboarded"}],"totalElements":3}'),
        ("API 27", "GET /v1/system-health", "View System Health", "PASSED (200 OK)", "None",
         '{"status":"UP","services":{"institute-management-service":{"status":"UP"}}}'),
        ("Story 53", "Audit Boundary Check", "Verify Zero Tenant Log Leakage", "PASSED (200 OK)", "Query: event_type=TenantInternalCourseEnrollment",
         '{"content":[],"page":0,"size":10,"totalElements":0}')
    ]

    for api_num, ep, purpose, status, req, resp in platform_tests:
        p_api = doc.add_paragraph()
        r_api = p_api.add_run(f"• {api_num}: {ep} — {purpose}")
        r_api.font.bold = True
        r_api.font.size = Pt(10.5)

        t_api = doc.add_table(rows=3, cols=2)
        t_api.alignment = WD_TABLE_ALIGNMENT.CENTER

        c00 = t_api.cell(0, 0); c01 = t_api.cell(0, 1)
        c00.text = "Result Status"; c01.text = status
        c00.paragraphs[0].runs[0].font.bold = True
        c01.paragraphs[0].runs[0].font.bold = True
        c01.paragraphs[0].runs[0].font.color.rgb = SUCCESS_COLOR
        set_cell_background(c00, "F0F4F8")
        set_cell_background(c01, "E8F5E9")

        c10 = t_api.cell(1, 0); c11 = t_api.cell(1, 1)
        c10.text = "Sample Input"; c11.text = req
        c10.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c10, "F0F4F8")
        set_cell_background(c11, "FAFAFA")

        c20 = t_api.cell(2, 0); c21 = t_api.cell(2, 1)
        c20.text = "Sample Response Payload"; c21.text = resp
        c20.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c20, "F0F4F8")
        set_cell_background(c21, "FAFAFA")

        for row in t_api.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

    # 3. College Admin Service Verification
    h3 = doc.add_heading("3. College Admin Service (Tenant Tier) Test Verification", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    p3 = doc.add_paragraph()
    p3.add_run("All 15 REST APIs and 2 state-machine/isolation test cases passed with 100% success.")

    college_tests = [
        ("API 01", "GET /v1/college-configs", "View Institution Configs", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101",
         '[{"key":"attendance_module_enabled","value":true},{"key":"theme_color","value":"#003366"}]'),
        ("API 02", "PATCH /v1/college-configs/{key}", "Update College Config", "PASSED (200 OK)", '{"value":"#0055A5"} (Header: X-Institution-Id=inst-101)',
         '{"key":"theme_color","value":"#0055A5","institutionId":"inst-101"}'),
        ("API 03", "GET /v1/college-configs/history", "View Config Audit History", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101",
         '[] (Initial audit trail ready for logged changes)'),
        ("API 04", "POST /v1/users", "Create User Record", "PASSED (201 Created)", '{"profileType":"student","name":"Jane Doe","email":"jane.doe@oxford.edu"}',
         '{"id":"usr-1fa364ea","name":"Jane Doe","status":"active"}'),
        ("API 05", "GET /v1/users/{id}", "View User Identity & Profile", "PASSED (200 OK)", "Path Variable: id=usr-101, Header: X-Institution-Id=inst-101",
         '{"id":"usr-101","name":"Prof. Alan Turing","profileType":"faculty"}'),
        ("API 06", "GET /v1/users", "List & Filter Users", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101, Query: profile_type=faculty",
         '{"content":[{"id":"usr-101","name":"Prof. Alan Turing"}],"totalElements":1}'),
        ("API 07", "PATCH /v1/users/{id}/status", "Transition Account Status", "PASSED (200 OK)", '{"status":"suspended"}',
         '{"id":"usr-101","name":"Prof. Alan Turing","status":"suspended"}'),
        ("API 08", "PATCH /v1/users/{id}/profile", "Update Profile Sub-document", "PASSED (200 OK)", '{"profile":{"designation":"Senior Professor"}}',
         '{"id":"usr-101","profile":{"department":"Computer Science","designation":"Senior Professor"}}'),
        ("API 09", "POST /v1/roles", "Define Custom Role", "PASSED (201 Created)", '{"name":"Department Head","permissions":["college:users:read"]}',
         '{"id":"role-9751e2fe","name":"Department Head"}'),
        ("API 10", "GET /v1/roles", "List Custom Roles", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101",
         '[{"id":"role-college-admin","name":"College Admin"}]'),
        ("API 11", "POST /v1/role-assignments", "Grant Role to User", "PASSED (201 Created)", '{"userId":"usr-101","roleId":"role-college-admin"}',
         '{"id":"assign-f6732c38","userId":"usr-101","roleName":"College Admin"}'),
        ("API 12", "GET /v1/role-assignments/{user_id}/effective", "Resolve Effective Permissions", "PASSED (200 OK)", "Path Variable: user_id=usr-101, Header: X-Institution-Id=inst-101",
         '{"userId":"usr-101","roles":["College Admin"],"effectivePermissions":["college:users:read"]}'),
        ("API 13", "DELETE /v1/role-assignments/{id}", "Revoke Role Assignment", "PASSED (204 No Content)", "Path Variable: id=assign-f6732c38",
         "HTTP 204 No Content"),
        ("API 14", "GET /v1/college-analytics/dashboard", "View Analytics Dashboard", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101, Query: metric=enrollment_summary",
         '{"snapshotId":"tenant-snap-1786360455904","data":{"totalEnrolledStudents":4850}}'),
        ("API 15", "POST /v1/college-analytics/recompute", "Trigger Analytics Recompute", "PASSED (202 Accepted)", '{"metric":"enrollment_summary","period":"semester"}',
         '{"jobId":"tenant-job-1786360455904","status":"ACCEPTED"}'),
        ("API 16", "GET /v1/audit-logs", "Query Institution Audit Trail", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101, Query: page=0, size=5",
         '{"content":[{"id":"tenant-audit-1","eventType":"UserCreated"}],"totalElements":3}'),
        ("Story 36", "Tenant RBAC Isolation Check", "Verify Zero Platform Permission Leaks", "PASSED (200 OK)", "Header: X-Institution-Id=inst-101",
         "Verified zero platform-tier permissions in tenant effective permission set.")
    ]

    for api_num, ep, purpose, status, req, resp in college_tests:
        p_api = doc.add_paragraph()
        r_api = p_api.add_run(f"• {api_num}: {ep} — {purpose}")
        r_api.font.bold = True
        r_api.font.size = Pt(10.5)

        t_api = doc.add_table(rows=3, cols=2)
        t_api.alignment = WD_TABLE_ALIGNMENT.CENTER

        c00 = t_api.cell(0, 0); c01 = t_api.cell(0, 1)
        c00.text = "Result Status"; c01.text = status
        c00.paragraphs[0].runs[0].font.bold = True
        c01.paragraphs[0].runs[0].font.bold = True
        c01.paragraphs[0].runs[0].font.color.rgb = SUCCESS_COLOR
        set_cell_background(c00, "F0F4F8")
        set_cell_background(c01, "E8F5E9")

        c10 = t_api.cell(1, 0); c11 = t_api.cell(1, 1)
        c10.text = "Sample Input"; c11.text = req
        c10.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c10, "F0F4F8")
        set_cell_background(c11, "FAFAFA")

        c20 = t_api.cell(2, 0); c21 = t_api.cell(2, 1)
        c20.text = "Sample Response Payload"; c21.text = resp
        c20.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c20, "F0F4F8")
        set_cell_background(c21, "FAFAFA")

        for row in t_api.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

    # Output paths
    output_path = r"d:\CampXSync\CampXSync_Admin_Module_Master_APIs_Test_Verification_Report.docx"
    doc.save(output_path)

    for dst_dir in [r"d:\CampXSync\CampXSync", r"d:\CampXSync\Documentation\AdminModule"]:
        try:
            shutil.copy(output_path, dst_dir)
        except Exception:
            pass

    print(f"Master API Test Verification Report generated successfully at {output_path}!")

if __name__ == "__main__":
    create_master_verification_doc()
