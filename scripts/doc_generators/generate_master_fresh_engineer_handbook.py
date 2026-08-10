import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import shutil

def create_fresh_engineer_handbook():
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Theme Colors
    PRIMARY_COLOR = RGBColor(0, 51, 102)     # Deep Navy Blue
    SECONDARY_COLOR = RGBColor(70, 130, 180)  # Steel Blue
    ACCENT_COLOR = RGBColor(180, 50, 50)     # Crimson Red
    SUCCESS_COLOR = RGBColor(0, 128, 0)      # Forest Green

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CampXSync Microservices & API Gateway\nFresh Engineer Master Onboarding & Port Configuration Handbook")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("A Comprehensive, Zero-Assumption Technical Guide for Recent College Graduates\nAuthor: Lead Principal Microservices Architect — World Top 1% Engineering Standard")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()

    # 1. Welcome & Handbook Purpose
    h1 = doc.add_heading("1. Executive Welcome & Core Architectural Principles", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p_wel = doc.add_paragraph()
    p_wel.add_run(
        "Welcome to the CampXSync Engineering Team! As a recent college graduate joining our software organization, "
        "this handbook is designed specifically for you. It covers every single technical step required to set up your workstation, "
        "build shared dependencies, launch microservices on specific port numbers (8080, 8088, 8089), understand HTTP request routing, "
        "and test REST APIs using cURL and Postman. Nothing is left to guesswork."
    )

    t_rules = doc.add_table(rows=5, cols=2)
    t_rules.alignment = WD_TABLE_ALIGNMENT.CENTER
    rules_data = [
        ("Golden Rule 1: Zero Assumptions", "Every directory path, terminal command, and configuration property is explicitly written out for you."),
        ("Golden Rule 2: Explicit Port Binding", "Never guess a microservice port. Always verify port bindings in application.yml or via explicit CLI parameters."),
        ("Golden Rule 3: Single Gateway Ingress", "All client applications (Web, Mobile, Postman) call API Gateway on Port 8080. Never bypass the Gateway."),
        ("Golden Rule 4: Structured Logging", "Never use System.out.println(). Always use AppLogger and AuditLogger from com.campxsync:logger.")
    ]
    for i, (k, v) in enumerate(rules_data):
        c0 = t_rules.cell(i, 0); c1 = t_rules.cell(i, 1)
        c0.text = k; c1.text = v
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9.5)
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")

    doc.add_paragraph()

    # 2. Deep Dive: How Microservice Port Assignment Works in Spring Boot
    h2 = doc.add_heading("2. Deep Dive: Understanding Spring Boot Port Configuration", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    p_port_intro = doc.add_paragraph()
    p_port_intro.add_run(
        "In a Java Spring Boot microservice architecture, each service runs an embedded web server (Tomcat/Jetty) "
        "that binds to a specific TCP network port on host machine (localhost). The default port is configured inside "
        "src/main/resources/application.yml under the property 'server.port'."
    )

    t_ports = doc.add_table(rows=5, cols=4)
    t_ports.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_headers = ["Microservice Name", "Configured Port", "Configuration File Path", "Role in System"]
    for j, h_text in enumerate(p_headers):
        cell = t_ports.cell(0, j)
        cell.text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    port_rows = [
        ("api-gateway", "8080", "api-gateway/src/main/resources/application.yml", "Central reverse proxy, JWT validation, MDC tracing"),
        ("platform-admin-service", "8088", "services/platform-admin-service/src/main/resources/application.yml", "Platform Tier: Institutes, global configs, billing, platform RBAC"),
        ("college-admin-service", "8089", "services/college-admin-service/src/main/resources/application.yml", "Tenant Tier: College configs, student/faculty profiles, tenant RBAC"),
        ("logger (library)", "N/A", "logger/pom.xml", "Shared Maven library (No HTTP port)")
    ]

    for i, (name, port, path, role) in enumerate(port_rows, start=1):
        c0 = t_ports.cell(i, 0); c1 = t_ports.cell(i, 1); c2 = t_ports.cell(i, 2); c3 = t_ports.cell(i, 3)
        c0.text = name; c1.text = port; c2.text = path; c3.text = role
        bg = "F9FBFD" if i % 2 == 0 else "FFFFFF"
        for c in [c0, c1, c2, c3]:
            c.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_background(c, bg)

    doc.add_paragraph()

    # 3. 3 Methods to Specify and Override Port Numbers
    h3 = doc.add_heading("3. Three Methods to Specify or Override Port Numbers", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph(
        "As a developer, you can specify or override the HTTP port of any microservice using 3 different methods "
        "without modifying the source code. Below are detailed instructions for each method:"
    )

    methods = [
        ("Method 1: Using Maven Command Line Arguments (--server.port)",
         "When running via 'mvn spring-boot:run', pass the port argument:\n\n"
         "• For API Gateway (Port 8080):\n"
         "  mvn spring-boot:run -Dspring-boot.run.arguments=\"--server.port=8080\"\n\n"
         "• For Platform Admin Service (Port 8088):\n"
         "  mvn spring-boot:run -Dspring-boot.run.arguments=\"--server.port=8088\"\n\n"
         "• For College Admin Service (Port 8089):\n"
         "  mvn spring-boot:run -Dspring-boot.run.arguments=\"--server.port=8089\"",
         "This method instructs Spring Boot to override server.port at application startup."),

        ("Method 2: Using Executable JAR Files (java -jar with System Property)",
         "First, package the project into a runnable JAR file:\n"
         "  mvn clean package -DskipTests\n\n"
         "Then run the compiled JAR with explicit port system property (-Dserver.port):\n\n"
         "• For API Gateway:\n"
         "  java -Dserver.port=8080 -jar target/api-gateway-1.0.0-SNAPSHOT.jar\n\n"
         "• For Platform Admin Service:\n"
         "  java -Dserver.port=8088 -jar target/platform-admin-service-1.0.0-SNAPSHOT.jar\n\n"
         "• For College Admin Service:\n"
         "  java -Dserver.port=8089 -jar target/college-admin-service-1.0.0-SNAPSHOT.jar",
         "This is the production standard method used when deploying microservices to Docker containers or cloud VM servers."),

        ("Method 3: Using OS Environment Variables (SERVER_PORT)",
         "Spring Boot automatically recognizes the SERVER_PORT environment variable.\n\n"
         "• In Windows PowerShell:\n"
         "  $env:SERVER_PORT=\"8080\"\n"
         "  mvn spring-boot:run\n\n"
         "• In Windows Command Prompt (cmd):\n"
         "  set SERVER_PORT=8080\n"
         "  mvn spring-boot:run",
         "This method is ideal when setting up environment configurations in CI/CD pipelines or Kubernetes deployments.")
    ]

    for m_title, m_cmd, m_desc in methods:
        p_m = doc.add_paragraph()
        r_m = p_m.add_run(f"• {m_title}")
        r_m.font.bold = True
        r_m.font.size = Pt(10.5)

        t_m = doc.add_table(rows=2, cols=2)
        t_m.alignment = WD_TABLE_ALIGNMENT.CENTER

        c00 = t_m.cell(0, 0); c01 = t_m.cell(0, 1)
        c00.text = "Execution Commands"; c01.text = m_cmd
        c00.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c00, "F0F4F8"); set_cell_background(c01, "FAFAFA")

        c10 = t_m.cell(1, 0); c11 = t_m.cell(1, 1)
        c10.text = "How It Works"; c11.text = m_desc
        c10.paragraphs[0].runs[0].font.bold = True
        set_cell_background(c10, "F0F4F8"); set_cell_background(c11, "E8F5E9")

        for row in t_m.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)

        doc.add_paragraph()

    # 4. Step-by-Step First Day Local Execution Walkthrough
    h4 = doc.add_heading("4. Step-by-Step First Day Local Execution Walkthrough", level=1)
    h4.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("Open 3 separate PowerShell terminal windows side-by-side and follow this exact sequence:")

    seq_steps = [
        ("Step A: Build & Install Shared Logger Library (MANDATORY FIRST STEP)",
         "Directory: d:\\CampXSync\\CampXSync\\logger\nCommand: mvn clean install\n"
         "Expected Log: [INFO] BUILD SUCCESS\n"
         "Explanation: Installs com.campxsync:logger:1.0.0 into your local ~/.m2 repository."),

        ("Step B: Launch Platform Admin Service on Port 8088",
         "Directory: d:\\CampXSync\\CampXSync\\services\\platform-admin-service\n"
         "Command: mvn spring-boot:run -Dspring-boot.run.arguments=\"--server.port=8088\"\n"
         "Expected Log: Tomcat started on port(s): 8088 (http) with context path ''"),

        ("Step C: Launch College Admin Service on Port 8089",
         "Directory: d:\\CampXSync\\CampXSync\\services\\college-admin-service\n"
         "Command: mvn spring-boot:run -Dspring-boot.run.arguments=\"--server.port=8089\"\n"
         "Expected Log: Tomcat started on port(s): 8089 (http) with context path ''"),

        ("Step D: Launch Central API Gateway on Port 8080",
         "Directory: d:\\CampXSync\\CampXSync\\api-gateway\n"
         "Command: mvn spring-boot:run -Dspring-boot.run.arguments=\"--server.port=8080\"\n"
         "Expected Log: Tomcat started on port(s): 8080 (http)\n"
         "Expected Log: CampXSync API Gateway microservice started successfully on port 8080.")
    ]

    for title, details in seq_steps:
        p_s = doc.add_paragraph()
        r_s = p_s.add_run(f"• {title}")
        r_s.font.bold = True
        r_s.font.size = Pt(10.5)

        p_det = doc.add_paragraph()
        r_det = p_det.add_run(details)
        r_det.font.name = "Consolas"; r_det.font.size = Pt(8.5)

        doc.add_paragraph()

    # 5. Complete API Catalog & Port Routing Table
    h5 = doc.add_heading("5. API Gateway Port Routing Catalog (All 42 REST APIs)", level=1)
    h5.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("All API calls MUST target http://localhost:8080. Gateway routes requests automatically:")

    apis = [
        ("POST /v1/institutes", "8088 (Platform)", "Provision institute", "{\"name\":\"Stanford\",\"subdomain\":\"stanford\",\"planId\":\"plan-enterprise\"}"),
        ("GET /v1/institutes/{id}", "8088 (Platform)", "View institute", "Path: id=inst-101"),
        ("GET /v1/institutes", "8088 (Platform)", "List institutes", "Query: ?status=active"),
        ("PATCH /v1/institutes/{id}/status", "8088 (Platform)", "Status transition", "{\"status\":\"suspended\"}"),
        ("PATCH /v1/institutes/{id}", "8088 (Platform)", "Update details", "{\"name\":\"Stanford University\"}"),
        ("GET /v1/platform-configs", "8088 (Platform)", "List platform configs", "None"),
        ("PATCH /v1/platform-configs/{key}", "8088 (Platform)", "Update config key", "{\"value\": false}"),
        ("GET /v1/platform-configs/history", "8088 (Platform)", "Config audit history", "Query: ?key=mfa_required"),
        ("POST /v1/platform-roles", "8088 (Platform)", "Define platform role", "{\"name\":\"Support Ops\",\"permissions\":[\"platform:institutes:read\"]}"),
        ("GET /v1/platform-roles", "8088 (Platform)", "List platform roles", "None"),
        ("POST /v1/platform-role-assignments", "8088 (Platform)", "Grant staff role", "{\"staffId\":\"staff-99\",\"roleId\":\"role-super-admin\"}"),
        ("GET /v1/platform-role-assignments/{id}/effective", "8088 (Platform)", "Effective permissions", "Path: id=staff-99"),
        ("DELETE /v1/platform-role-assignments/{id}", "8088 (Platform)", "Revoke role grant", "Path: id=assign-101"),
        ("POST /v1/policies", "8088 (Platform)", "Create governance policy", "{\"name\":\"GDPR Retention\",\"type\":\"retention\",\"appliesTo\":\"all\"}"),
        ("GET /v1/policies", "8088 (Platform)", "List policies", "Query: ?applies_to=all"),
        ("PATCH /v1/policies/{id}", "8088 (Platform)", "Update policy rule", "{\"name\":\"Updated Retention Policy\"}"),
        ("DELETE /v1/policies/{id}", "8088 (Platform)", "Retire policy", "Path: id=pol-101"),
        ("GET /v1/billing-accounts/{id}", "8088 (Platform)", "View billing status", "Path: id=inst-101"),
        ("POST /v1/billing-accounts/{id}/change-plan", "8088 (Platform)", "Change subscription plan", "{\"newPlanId\":\"plan-enterprise-v2\"}"),
        ("POST /v1/billing-accounts/{id}/charge", "8088 (Platform)", "Settle manual charge", "{\"amount\":1250.00,\"description\":\"Overage\"}"),
        ("GET /v1/analytics/snapshots", "8088 (Platform)", "Platform rollup metrics", "Query: ?metric=active_institutions_count"),
        ("POST /v1/analytics/snapshots/recompute", "8088 (Platform)", "Recompute analytics", "{\"metric\":\"active_institutions_count\",\"period\":\"monthly\"}"),
        ("POST /v1/compliance-checks/run", "8088 (Platform)", "Trigger policy audit", "{\"institutionId\":\"inst-101\",\"policyIds\":[\"pol-101\"]}"),
        ("GET /v1/compliance-checks/{id}", "8088 (Platform)", "View compliance results", "Path: id=inst-101"),
        ("GET /v1/compliance-checks", "8088 (Platform)", "List non-compliant inst", "Query: ?flagged=true"),
        ("GET /v1/platform-audit-logs", "8088 (Platform)", "Query platform audit log", "Query: ?page=0&size=5"),
        ("GET /v1/system-health", "8088 (Platform)", "System health check", "None"),
        ("GET /v1/college-configs", "8089 (College)", "View college configs", "Header: X-Institution-Id: inst-101"),
        ("PATCH /v1/college-configs/{key}", "8089 (College)", "Update college config", "{\"value\":\"#0055A5\"} (Header: X-Institution-Id: inst-101)"),
        ("GET /v1/college-configs/history", "8089 (College)", "Config audit history", "Header: X-Institution-Id: inst-101"),
        ("POST /v1/users", "8089 (College)", "Create user profile", "{\"profileType\":\"student\",\"name\":\"Jane\",\"email\":\"jane@oxford.edu\"}"),
        ("GET /v1/users/{id}", "8089 (College)", "View user profile", "Header: X-Institution-Id: inst-101, Path: id=usr-101"),
        ("GET /v1/users", "8089 (College)", "List tenant users", "Header: X-Institution-Id: inst-101, Query: ?profile_type=faculty"),
        ("PATCH /v1/users/{id}/status", "8089 (College)", "User status transition", "{\"status\":\"suspended\"} (Header: X-Institution-Id: inst-101)"),
        ("PATCH /v1/users/{id}/profile", "8089 (College)", "Update profile subdoc", "{\"profile\":{\"designation\":\"Senior Professor\"}}"),
        ("POST /v1/roles", "8089 (College)", "Define custom role", "{\"name\":\"Dept Head\",\"permissions\":[\"college:users:read\"]}"),
        ("GET /v1/roles", "8089 (College)", "List tenant roles", "Header: X-Institution-Id: inst-101"),
        ("POST /v1/role-assignments", "8089 (College)", "Grant role to user", "{\"userId\":\"usr-101\",\"roleId\":\"role-college-admin\"}"),
        ("GET /v1/role-assignments/{id}/effective", "8089 (College)", "Effective permissions", "Header: X-Institution-Id: inst-101, Path: id=usr-101"),
        ("DELETE /v1/role-assignments/{id}", "8089 (College)", "Revoke tenant role grant", "Header: X-Institution-Id: inst-101, Path: id=assign-101"),
        ("GET /v1/college-analytics/dashboard", "8089 (College)", "View tenant dashboard", "Header: X-Institution-Id: inst-101, Query: ?metric=enrollment_summary"),
        ("POST /v1/college-analytics/recompute", "8089 (College)", "Recompute tenant dashboard", "{\"metric\":\"enrollment_summary\",\"period\":\"semester\"}"),
        ("GET /v1/audit-logs", "8089 (College)", "Query tenant audit log", "Header: X-Institution-Id: inst-101, Query: ?page=0&size=5")
    ]

    t_a = doc.add_table(rows=len(apis) + 1, cols=4)
    t_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    a_headers = ["Gateway Path (Port 8080)", "Target Port", "Functional Purpose", "Sample Request Details"]
    for j, h_text in enumerate(a_headers):
        cell = t_a.cell(0, j)
        cell.text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (path, port, purpose, details) in enumerate(apis, start=1):
        c0 = t_a.cell(i, 0); c1 = t_a.cell(i, 1); c2 = t_a.cell(i, 2); c3 = t_a.cell(i, 3)
        c0.text = path; c1.text = port; c2.text = purpose; c3.text = details
        bg = "F9FBFD" if i % 2 == 0 else "FFFFFF"
        for c in [c0, c1, c2, c3]:
            c.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_background(c, bg)

    doc.add_paragraph()

    # 6. Port Conflict Management & Troubleshooting
    h6 = doc.add_heading("6. Port Conflict Management & Emergency Playbook", level=1)
    h6.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph("If a microservice fails to start due to a port conflict, use this exact troubleshooting procedure:")

    trouble = [
        ("Problem: 'Port 8080 (or 8088 / 8089) is already in use'",
         "Step 1: Open PowerShell as Administrator.\n"
         "Step 2: Run netstat to find process ID (PID):\n"
         "        netstat -ano | findstr :8080\n"
         "Step 3: Kill the conflicting process:\n"
         "        taskkill /F /PID <PID>\n"
         "Step 4: Re-run your microservice command."),

        ("Problem: 'HTTP 502 Bad Gateway' when calling Port 8080",
         "Cause: API Gateway (8080) is running, but the downstream microservice (8088 or 8089) is not running.\n"
         "Fix: Open Terminal 1 or Terminal 2 and start the corresponding service using 'mvn spring-boot:run'."),

        ("Problem: Maven error 'Could not resolve dependencies com.campxsync:logger:1.0.0'",
         "Cause: You forgot to run Step A (building the logger library).\n"
         "Fix: Run `cd d:\\CampXSync\\CampXSync\\logger && mvn clean install`.")
    ]

    for p_title, p_body in trouble:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(f"• {p_title}")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = ACCENT_COLOR

        p_b = doc.add_paragraph()
        r_b = p_b.add_run(p_body)
        r_b.font.name = "Consolas"; r_b.font.size = Pt(8.5)

        doc.add_paragraph()

    # Save Master Handbook
    output_path = r"d:\CampXSync\CampXSync_Complete_Junior_Engineer_Onboarding_Handbook.docx"
    try:
        doc.save(output_path)
    except PermissionError:
        output_path = r"d:\CampXSync\CampXSync_Complete_Junior_Engineer_Onboarding_Handbook_v2.0.docx"
        doc.save(output_path)

    for dst_dir in [r"d:\CampXSync\CampXSync", r"d:\CampXSync\Documentation\AdminModule"]:
        try:
            shutil.copy(output_path, dst_dir)
        except Exception:
            pass

    print(f"Enriched Junior Engineer Onboarding Handbook generated successfully at {output_path}!")

if __name__ == "__main__":
    create_fresh_engineer_handbook()
